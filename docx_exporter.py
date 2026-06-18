from io import BytesIO
from datetime import date
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from config import WORD_FONT, DOCX_CELL_COLORS, SPEAKER_COLORS


def _set_cell_bg(cell, hex_color: str) -> None:
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color.upper())
    tcPr.append(shd)


def _set_run_font(run, font_name: str = WORD_FONT, size_pt: int = 10) -> None:
    run.font.name = font_name
    run.font.size = Pt(size_pt)
    try:
        rPr = run._element.get_or_add_rPr()
        rPr.rFonts.set(qn("w:eastAsia"), font_name)
    except Exception:
        pass


def _add_heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.name = WORD_FONT
        try:
            rPr = run._element.get_or_add_rPr()
            rPr.rFonts.set(qn("w:eastAsia"), WORD_FONT)
        except Exception:
            pass


def _add_table_with_headers(doc: Document, headers: list[str], col_widths_cm: list[float]) -> object:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        _set_cell_bg(hdr_cells[i], "D5D8DC")
        run = hdr_cells[i].paragraphs[0].runs[0]
        run.bold = True
        _set_run_font(run, size_pt=10)
    for i, width in enumerate(col_widths_cm):
        for cell in table.column_cells(i):
            cell.width = Cm(width)
    return table


def _add_cell_text(cell, text: str, font_size: int = 10, bold: bool = False, color_hex: str = None) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.bold = bold
    _set_run_font(run, size_pt=font_size)
    if color_hex:
        r, g, b = int(color_hex[0:2], 16), int(color_hex[2:4], 16), int(color_hex[4:6], 16)
        run.font.color.rgb = RGBColor(r, g, b)


def _add_paragraph_colored_bg(doc: Document, text: str, bg_hex: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    _set_run_font(run)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), bg_hex.upper())
    pPr.append(shd)


def _build_lyrics_section(doc: Document, lines: list[dict]) -> None:
    _add_heading(doc, "📖 전체 가사 — 원문·음차·해석", level=2)
    headers = ["📄 일본어 원문", "🔤 한국어 음차", "🇰🇷 한국어 해석"]
    col_widths = [7.0, 5.5, 5.5]
    table = _add_table_with_headers(doc, headers, col_widths)

    for line in lines:
        if line.get("is_section_break"):
            row = table.add_row()
            for cell in row.cells:
                _set_cell_bg(cell, "F2F3F4")
        else:
            row = table.add_row()
            _add_cell_text(row.cells[0], line.get("japanese", ""), font_size=12)
            _add_cell_text(row.cells[1], line.get("korean_phonetic", ""), font_size=10, color_hex="626567")
            _add_cell_text(row.cells[2], line.get("korean_translation", ""), font_size=10)


def _build_vocabulary_section(doc: Document, vocabulary: list[dict]) -> None:
    _add_heading(doc, "📚 단어장", level=2)
    headers = ["단어", "읽기", "한국어 음차", "뜻", "JLPT", "가사 예문"]
    col_widths = [2.5, 3.0, 3.0, 3.5, 1.5, 4.5]
    table = _add_table_with_headers(doc, headers, col_widths)

    for v in vocabulary:
        level = v.get("jlpt_level", "unknown")
        cell_color = DOCX_CELL_COLORS.get(level, DOCX_CELL_COLORS["unknown"])

        meaning = v.get("meaning_korean", "")
        if v.get("is_polysemous"):
            meaning += " ⚡"
        note = v.get("colloquial_note")
        if note:
            meaning += f"\n⚠️ {note}"

        collocations = v.get("collocations", [])
        example = f"예문 줄 #{v.get('example_line_number', '')}"
        if collocations:
            example += "\n콜로: " + " / ".join(collocations[:3])

        row = table.add_row()
        _add_cell_text(row.cells[0], v.get("word", ""), bold=True)
        _set_cell_bg(row.cells[0], cell_color)
        _add_cell_text(row.cells[1], v.get("reading", ""))
        _add_cell_text(row.cells[2], v.get("korean_phonetic", ""))
        _add_cell_text(row.cells[3], meaning)
        _add_cell_text(row.cells[4], level)
        _set_cell_bg(row.cells[4], cell_color)
        _add_cell_text(row.cells[5], example)


def _build_grammar_section(doc: Document, grammar_points: list[dict]) -> None:
    _add_heading(doc, "📐 문법정리", level=2)
    headers = ["패턴", "한국어 음차", "기능 태그", "설명", "JLPT", "예문 / 주의"]
    col_widths = [3.5, 3.5, 2.5, 3.5, 1.5, 4.0]
    table = _add_table_with_headers(doc, headers, col_widths)

    for g in grammar_points:
        level = g.get("jlpt_level", "unknown")
        cell_color = DOCX_CELL_COLORS.get(level, DOCX_CELL_COLORS["unknown"])
        tags = " ".join([f"[{t}]" for t in g.get("function_tags", [])])

        extra = g.get("example_japanese", "")
        confusion = g.get("confusion_note")
        if confusion:
            extra += f"\n📌 {confusion}"
        standard = g.get("standard_form")
        if standard:
            extra += f"\n→ 표준형: {standard}"

        row = table.add_row()
        _add_cell_text(row.cells[0], g.get("pattern", ""), bold=True)
        _set_cell_bg(row.cells[0], cell_color)
        _add_cell_text(row.cells[1], g.get("pattern_korean_phonetic", ""))
        _add_cell_text(row.cells[2], tags)
        _add_cell_text(row.cells[3], g.get("explanation_korean", ""))
        _add_cell_text(row.cells[4], level)
        _set_cell_bg(row.cells[4], cell_color)
        _add_cell_text(row.cells[5], extra)


def _build_kanji_section(doc: Document, kanji_list: list[dict]) -> None:
    _add_heading(doc, "🈶 한자 목록 (JLPT N1~N5)", level=2)
    headers = ["한자", "한국식 음", "음독(音読み)", "훈독(訓読み)", "뜻 / 예시 단어", "JLPT / ★"]
    col_widths = [2.0, 2.5, 3.0, 3.0, 5.0, 2.5]
    table = _add_table_with_headers(doc, headers, col_widths)

    for k in kanji_list:
        level = k.get("jlpt_level", "N5")
        cell_color = DOCX_CELL_COLORS.get(level, DOCX_CELL_COLORS["unknown"])
        star = " ★" if k.get("is_key_kanji") else ""
        example_words = " / ".join(k.get("example_words", [])[:3])
        meaning_extra = k.get("meaning_korean", "")
        if example_words:
            meaning_extra += f"\n예시: {example_words}"

        row = table.add_row()
        _add_cell_text(row.cells[0], k.get("kanji", ""), bold=True, font_size=14)
        _set_cell_bg(row.cells[0], cell_color)
        _add_cell_text(row.cells[1], k.get("korean_reading", ""))
        _add_cell_text(row.cells[2], k.get("reading_on", ""))
        _add_cell_text(row.cells[3], k.get("reading_kun", ""))
        _add_cell_text(row.cells[4], meaning_extra)
        _add_cell_text(row.cells[5], f"{level}{star}")
        _set_cell_bg(row.cells[5], cell_color)


def _build_conversation_section(doc: Document, conversation_practice: list[dict]) -> None:
    _add_heading(doc, "💬 간단 회화 실습", level=2)
    for i, practice in enumerate(conversation_practice, 1):
        p = doc.add_heading(f"📍 상황 {i}: {practice.get('situation', '')}", level=3)
        for run in p.runs:
            run.font.name = WORD_FONT
            try:
                rPr = run._element.get_or_add_rPr()
                rPr.rFonts.set(qn("w:eastAsia"), WORD_FONT)
            except Exception:
                pass

        for turn in practice.get("dialogue", []):
            speaker = turn.get("speaker", "A")
            bg = SPEAKER_COLORS.get(speaker, "#F2F3F4").lstrip("#")

            lines = [
                f"{speaker}: {turn.get('japanese', '')}",
                f"  ({turn.get('korean_phonetic', '')})",
                f"  {turn.get('korean_translation', '')}",
                f"  ✨ {', '.join(turn.get('used_items', []))}",
            ]
            text = "\n".join(lines)
            _add_paragraph_colored_bg(doc, text, bg)

        doc.add_paragraph()


def export_to_docx(analysis: dict, song_title: str) -> bytes:
    doc = Document()

    for style in doc.styles:
        try:
            if style.font:
                style.font.name = WORD_FONT
                try:
                    rPr = style.element.get_or_add_rPr()
                    rPr.rFonts.set(qn("w:eastAsia"), WORD_FONT)
                except Exception:
                    pass
        except AttributeError:
            pass

    song_info = analysis.get("song_info", {})
    artist = song_info.get("artist", "")
    level_est = song_info.get("overall_jlpt_estimate", "")
    today = date.today().strftime("%Y-%m-%d")

    title_para = doc.add_paragraph()
    title_run = title_para.add_run(f"{song_title} — {artist}")
    title_run.bold = True
    title_run.font.size = Pt(18)
    _set_run_font(title_run, size_pt=18)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    sub_para = doc.add_paragraph()
    sub_run = sub_para.add_run(f"JLPT 추정 레벨: {level_est}  |  생성일: {today}")
    sub_run.font.size = Pt(10)
    _set_run_font(sub_run, size_pt=10)
    sub_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    _build_lyrics_section(doc, analysis.get("lines", []))
    doc.add_page_break()

    _build_vocabulary_section(doc, analysis.get("vocabulary", []))
    doc.add_page_break()

    _build_grammar_section(doc, analysis.get("grammar_points", []))
    doc.add_paragraph()

    _build_kanji_section(doc, analysis.get("kanji_list", []))
    doc.add_page_break()

    _build_conversation_section(doc, analysis.get("conversation_practice", []))

    _add_heading(doc, "📊 JLPT 분포", level=2)
    distribution = analysis.get("jlpt_distribution", {})
    dist_lines = [
        f"N5: {distribution.get('N5', 0)}개  N4: {distribution.get('N4', 0)}개  "
        f"N3: {distribution.get('N3', 0)}개  N2: {distribution.get('N2', 0)}개  "
        f"N1: {distribution.get('N1', 0)}개  미분류: {distribution.get('unknown', 0)}개",
        f"전체 레벨 추정: {level_est}",
    ]
    for line in dist_lines:
        p = doc.add_paragraph(line)
        for run in p.runs:
            _set_run_font(run)

    _add_heading(doc, "🎯 시험 학습 포인트", level=2)
    summary_para = doc.add_paragraph(analysis.get("study_summary", ""))
    for run in summary_para.runs:
        _set_run_font(run)

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()
